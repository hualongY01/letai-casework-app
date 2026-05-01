from __future__ import annotations

import json
from pathlib import Path
from uuid import uuid4

import fitz
import openpyxl
from docx import Document
from sqlmodel import Session, select

from letai_factbase.models import DocumentChunk, EvidenceSpan, FactCandidate, FactCard, Source


SUPPORTED_TEXT_SUFFIXES = {".txt", ".md", ".markdown"}
SUPPORTED_PDF_SUFFIXES = {".pdf"}
SUPPORTED_DOCX_SUFFIXES = {".docx"}
SUPPORTED_XLSX_SUFFIXES = {".xlsx"}


def _split_text(text: str, max_chars: int = 1800) -> list[tuple[str, int, int]]:
    paragraphs = text.splitlines(keepends=True)
    chunks: list[tuple[str, int, int]] = []
    current: list[str] = []
    start_line = 1
    line_no = 1

    for paragraph in paragraphs:
        if not current:
            start_line = line_no
        current.append(paragraph)
        current_text = "".join(current).strip()
        if len(current_text) >= max_chars:
            chunks.append((current_text, start_line, line_no))
            current = []
        line_no += 1

    tail = "".join(current).strip()
    if tail:
        chunks.append((tail, start_line, max(start_line, line_no - 1)))
    return chunks


def _create_chunk_with_evidence(
    session: Session,
    source_uid: str,
    text: str,
    locator: dict,
    page_number: int | None = None,
    sheet_name: str | None = None,
    row_start: int | None = None,
    row_end: int | None = None,
    column_start: str | None = None,
    column_end: str | None = None,
) -> DocumentChunk:
    chunk_uid = f"CHUNK-{uuid4().hex[:12]}"
    evidence_uid = f"EVID-{uuid4().hex[:12]}"
    chunk = DocumentChunk(
        chunk_uid=chunk_uid,
        source_uid=source_uid,
        text=text,
        page_start=page_number,
        page_end=page_number,
        sheet_name=sheet_name,
        row_start=row_start,
        row_end=row_end,
        column_start=column_start,
        column_end=column_end,
    )
    evidence = EvidenceSpan(
        evidence_uid=evidence_uid,
        source_uid=source_uid,
        chunk_uid=chunk_uid,
        excerpt=text,
        locator_json=json.dumps(locator, ensure_ascii=False),
        confidence=1.0,
    )
    session.add(chunk)
    session.add(evidence)
    return chunk


def _get_source(session: Session, source_uid: str) -> Source:
    source = session.exec(select(Source).where(Source.source_uid == source_uid)).first()
    if source is None:
        raise ValueError(f"Source not found: {source_uid}")
    return source


def parse_source_file(session: Session, source_uid: str, force: bool = False) -> list[DocumentChunk]:
    source = _get_source(session, source_uid)
    archive_path = Path(source.archive_path)
    suffix = archive_path.suffix.lower()
    existing_chunks = _get_existing_chunks(session, source_uid)
    if existing_chunks and not force:
        return existing_chunks
    if existing_chunks and force:
        _delete_existing_parse_outputs(session, source_uid, existing_chunks)

    if suffix in SUPPORTED_TEXT_SUFFIXES:
        created_chunks = _parse_text_file(session, source_uid, archive_path)
    elif suffix in SUPPORTED_PDF_SUFFIXES:
        created_chunks = _parse_pdf_file(session, source_uid, archive_path)
    elif suffix in SUPPORTED_DOCX_SUFFIXES:
        created_chunks = _parse_docx_file(session, source_uid, archive_path)
    elif suffix in SUPPORTED_XLSX_SUFFIXES:
        created_chunks = _parse_xlsx_file(session, source_uid, archive_path)
    else:
        raise ValueError(f"Unsupported source type: {archive_path.suffix}")

    session.commit()
    for chunk in created_chunks:
        session.refresh(chunk)
    return created_chunks


def _get_existing_chunks(session: Session, source_uid: str) -> list[DocumentChunk]:
    return list(
        session.exec(select(DocumentChunk).where(DocumentChunk.source_uid == source_uid)).all()
    )


def _delete_existing_parse_outputs(
    session: Session,
    source_uid: str,
    chunks: list[DocumentChunk],
) -> None:
    chunk_uids = [chunk.chunk_uid for chunk in chunks]
    evidence = list(
        session.exec(
            select(EvidenceSpan).where(
                EvidenceSpan.source_uid == source_uid,
                EvidenceSpan.chunk_uid.in_(chunk_uids),
            )
        ).all()
    )
    evidence_uids = [item.evidence_uid for item in evidence]
    if evidence_uids and _has_evidence_dependencies(session, evidence_uids):
        raise ValueError("Cannot force re-parse: existing evidence is referenced by candidates or FCs")

    for item in evidence:
        session.delete(item)
    for chunk in chunks:
        session.delete(chunk)
    session.flush()


def _has_evidence_dependencies(session: Session, evidence_uids: list[str]) -> bool:
    candidate = session.exec(
        select(FactCandidate).where(FactCandidate.evidence_uid.in_(evidence_uids))
    ).first()
    if candidate is not None:
        return True
    fact = session.exec(select(FactCard).where(FactCard.evidence_uid.in_(evidence_uids))).first()
    return fact is not None


def _parse_text_file(session: Session, source_uid: str, archive_path: Path) -> list[DocumentChunk]:
    text = archive_path.read_text(encoding="utf-8")
    parsed_chunks = _split_text(text)
    created_chunks: list[DocumentChunk] = []

    for index, (chunk_text, start_line, end_line) in enumerate(parsed_chunks, start=1):
        locator = {
            "source_uid": source_uid,
            "line_start": start_line,
            "line_end": end_line,
            "chunk_index": index,
            "parser": "text",
        }
        created_chunks.append(
            _create_chunk_with_evidence(
                session=session,
                source_uid=source_uid,
                text=chunk_text,
                locator=locator,
            )
        )
    return created_chunks


def _parse_pdf_file(session: Session, source_uid: str, archive_path: Path) -> list[DocumentChunk]:
    created_chunks: list[DocumentChunk] = []
    with fitz.open(archive_path) as document:
        for page_index, page in enumerate(document, start=1):
            page_text = page.get_text("text").strip()
            if not page_text:
                continue
            for chunk_index, (chunk_text, line_start, line_end) in enumerate(
                _split_text(page_text), start=1
            ):
                locator = {
                    "source_uid": source_uid,
                    "page_number": page_index,
                    "line_start": line_start,
                    "line_end": line_end,
                    "chunk_index": chunk_index,
                    "parser": "pdf_text",
                }
                created_chunks.append(
                    _create_chunk_with_evidence(
                        session=session,
                        source_uid=source_uid,
                        text=chunk_text,
                        locator=locator,
                        page_number=page_index,
                    )
                )
    return created_chunks


def _parse_docx_file(session: Session, source_uid: str, archive_path: Path) -> list[DocumentChunk]:
    document = Document(str(archive_path))
    blocks: list[tuple[str, dict]] = []

    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            blocks.append(
                (
                    text,
                    {
                        "source_uid": source_uid,
                        "paragraph_index": index,
                        "parser": "docx_paragraph",
                    },
                )
            )

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                text = " | ".join(values)
                blocks.append(
                    (
                        text,
                        {
                            "source_uid": source_uid,
                            "table_index": table_index,
                            "row_index": row_index,
                            "parser": "docx_table",
                        },
                    )
                )

    created_chunks: list[DocumentChunk] = []
    for block_index, (text, locator) in enumerate(blocks, start=1):
        locator["chunk_index"] = block_index
        created_chunks.append(
            _create_chunk_with_evidence(
                session=session,
                source_uid=source_uid,
                text=text,
                locator=locator,
            )
        )
    return created_chunks


def _parse_xlsx_file(session: Session, source_uid: str, archive_path: Path) -> list[DocumentChunk]:
    workbook = openpyxl.load_workbook(archive_path, data_only=True, read_only=True)
    created_chunks: list[DocumentChunk] = []

    for sheet in workbook.worksheets:
        headers: list[str] = []
        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = ["" if value is None else str(value).strip() for value in row]
            if not any(values):
                continue
            if not headers:
                headers = [value or f"Column {idx}" for idx, value in enumerate(values, start=1)]
                text = "Header: " + " | ".join(headers)
            else:
                pairs = []
                for idx, value in enumerate(values, start=1):
                    if value:
                        header = headers[idx - 1] if idx - 1 < len(headers) else f"Column {idx}"
                        pairs.append(f"{header}: {value}")
                text = "; ".join(pairs)

            locator = {
                "source_uid": source_uid,
                "sheet_name": sheet.title,
                "row_start": row_index,
                "row_end": row_index,
                "parser": "xlsx_row",
            }
            created_chunks.append(
                _create_chunk_with_evidence(
                    session=session,
                    source_uid=source_uid,
                    text=text,
                    locator=locator,
                    sheet_name=sheet.title,
                    row_start=row_index,
                    row_end=row_index,
                )
            )

    workbook.close()
    return created_chunks
