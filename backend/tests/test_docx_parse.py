from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from letai_factbase.db.session import init_db
from letai_factbase.main import app


def _make_docx(path: Path) -> bytes:
    document = Document()
    document.add_paragraph("DOCX 测试材料：用于验证段落解析。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "内容"
    table.cell(1, 0).text = "用途"
    table.cell(1, 1).text = "验证表格解析"
    document.save(path)
    return path.read_bytes()


def test_docx_parse_creates_chunks_for_paragraphs_and_tables() -> None:
    init_db()
    client = TestClient(app)
    docx_bytes = _make_docx(Path("/private/tmp/letai-docx-parse-test.docx"))

    import_response = client.post(
        "/api/sources/import",
        files={
            "file": (
                "docx-parse-test.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "title": "DOCX Parse Test Source",
            "imported_by": "pytest",
            "authority_level": "99",
        },
    )
    assert import_response.status_code == 200
    source_uid = import_response.json()["source_uid"]

    parse_response = client.post(f"/api/sources/{source_uid}/parse")

    assert parse_response.status_code == 200
    assert parse_response.json()["chunks_created"] >= 2

    chunks_response = client.get("/api/chunks")
    chunks = [chunk for chunk in chunks_response.json() if chunk["source_uid"] == source_uid]
    assert any("DOCX 测试材料" in chunk["text"] for chunk in chunks)
    assert any("验证表格解析" in chunk["text"] for chunk in chunks)
